import os
import pdfplumber
import docx
import string
import nltk
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
import re
from difflib import SequenceMatcher
import numpy as np
from collections import Counter
import json

# Download required NLTK data
try:
    nltk.download('stopwords', quiet=True)
    nltk.download('punkt', quiet=True)
except:
    pass

class EnhancedMERNResumeMatcher:
    def __init__(self):
        # Advanced TF-IDF configuration for better semantic understanding
        self.vectorizer = TfidfVectorizer(
            max_features=25000,
            stop_words=None,
            ngram_range=(1, 4),  # Extended n-grams for better phrase matching
            min_df=1,
            max_df=0.95,
            sublinear_tf=True,
            token_pattern=r'\b\w+(?:\.\w+)*\b',  # Better pattern for tech terms
            lowercase=True,
            use_idf=True
        )
        self.stemmer = PorterStemmer()
        
        # Enhanced skill categories with weighted importance for MERN stack
        self.skill_categories = {
            # Core MERN Stack (High Priority)
            'mongodb': {
                'terms': ['mongodb', 'mongo', 'mongoose', 'atlas', 'nosql', 'document database', 'bson'],
                'weight': 2.5,
                'category': 'Database'
            },
            'express': {
                'terms': ['express', 'expressjs', 'express.js', 'express server', 'middleware'],
                'weight': 2.5,
                'category': 'Backend Framework'
            },
            'react': {
                'terms': ['react', 'reactjs', 'react.js', 'jsx', 'react hooks', 'react router', 'redux', 
                         'context api', 'react native', 'next.js', 'nextjs', 'gatsby'],
                'weight': 2.5,
                'category': 'Frontend Framework'
            },
            'nodejs': {
                'terms': ['nodejs', 'node.js', 'node', 'npm', 'yarn', 'event loop', 'async/await'],
                'weight': 2.5,
                'category': 'Runtime Environment'
            },
            
            # Essential JavaScript/TypeScript (High Priority)
            'javascript': {
                'terms': ['javascript', 'js', 'ecmascript', 'es6', 'es2015', 'es2017', 'es2020', 
                         'vanilla js', 'modern javascript', 'async', 'promises', 'closures'],
                'weight': 2.0,
                'category': 'Programming Language'
            },
            'typescript': {
                'terms': ['typescript', 'ts', 'type annotations', 'interfaces', 'generics'],
                'weight': 1.8,
                'category': 'Programming Language'
            },
            
            # Frontend Technologies (Medium-High Priority)
            'html_css': {
                'terms': ['html', 'html5', 'css', 'css3', 'sass', 'scss', 'less', 'responsive design',
                         'flexbox', 'grid', 'bootstrap', 'tailwind', 'material-ui', 'styled-components'],
                'weight': 1.5,
                'category': 'Frontend Technologies'
            },
            
            # State Management (Medium-High Priority)
            'state_management': {
                'terms': ['redux', 'redux toolkit', 'context api', 'zustand', 'recoil', 'mobx',
                         'state management', 'global state'],
                'weight': 1.8,
                'category': 'State Management'
            },
            
            # Development Tools (Medium Priority)
            'development_tools': {
                'terms': ['webpack', 'babel', 'vite', 'parcel', 'eslint', 'prettier', 'git', 'github',
                         'gitlab', 'version control', 'npm', 'yarn', 'package manager'],
                'weight': 1.3,
                'category': 'Development Tools'
            },
            
            # Testing (Medium Priority)
            'testing': {
                'terms': ['jest', 'enzyme', 'react testing library', 'cypress', 'selenium', 'mocha',
                         'chai', 'unit testing', 'integration testing', 'e2e testing', 'tdd', 'bdd'],
                'weight': 1.4,
                'category': 'Testing'
            },
            
            # Cloud & Deployment (Medium Priority)
            'cloud_deployment': {
                'terms': ['aws', 'azure', 'gcp', 'heroku', 'netlify', 'vercel', 'docker', 'kubernetes',
                         'k8s', 'ci/cd', 'jenkins', 'github actions', 'gitlab ci'],
                'weight': 1.3,
                'category': 'Cloud & Deployment'
            },
            
            # API & Communication (Medium Priority)
            'api_communication': {
                'terms': ['rest', 'restful', 'api', 'graphql', 'apollo', 'axios', 'fetch', 'websockets',
                         'socket.io', 'jwt', 'authentication', 'authorization'],
                'weight': 1.6,
                'category': 'API & Communication'
            },
            
            # Additional Databases (Lower Priority for MERN but relevant)
            'other_databases': {
                'terms': ['postgresql', 'mysql', 'redis', 'elasticsearch', 'firebase', 'firestore'],
                'weight': 1.0,
                'category': 'Additional Databases'
            },
            
            # Soft Skills & Methodologies (Lower Priority but Important)
            'methodologies': {
                'terms': ['agile', 'scrum', 'kanban', 'sprint', 'jira', 'trello', 'collaboration',
                         'team work', 'problem solving', 'debugging'],
                'weight': 0.8,
                'category': 'Methodologies'
            }
        }

    def extract_text(self, file):
        """Enhanced text extraction with better error handling"""
        if hasattr(file, 'name'):
            _, ext = os.path.splitext(file.name)
        else:
            raise ValueError("Invalid file object")
            
        ext = ext.lower().strip()
        text = ""

        try:
            if ext == ".pdf":
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + " "
                        
                        # Extract text from tables if present
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    text += " ".join([cell for cell in row if cell]) + " "
                                    
            elif ext == ".docx":
                doc = docx.Document(file)
                # Extract from paragraphs
                text = " ".join([para.text for para in doc.paragraphs])
                
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += " " + cell.text
                            
            elif ext == ".txt":
                text = file.read().decode("utf-8")
            else:
                raise ValueError(f"Unsupported file format: '{ext}'. Use PDF, DOCX, or TXT.")
                
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"Error extracting text from {ext} file: {str(e)}")

    def advanced_text_preprocessing(self, text):
        """Advanced preprocessing specifically optimized for tech resumes"""
        if not text:
            return ""
        
        # Normalize text
        text = text.lower()
        
        # Handle common tech abbreviations and variations
        tech_normalizations = {
            r'\bnode\.?js\b': 'nodejs',
            r'\breact\.?js\b': 'react',
            r'\bexpress\.?js\b': 'express',
            r'\bnext\.?js\b': 'nextjs',
            r'\bmongo\s*db\b': 'mongodb',
            r'\btype\s*script\b': 'typescript',
            r'\bj\s*s\b': 'javascript',
            r'\bhtml\s*5?\b': 'html',
            r'\bcss\s*3?\b': 'css',
            r'\brest\s*ful?\b': 'restful',
            r'\bapi\s*s?\b': 'api',
            r'\bci\s*/?\s*cd\b': 'cicd',
            r'\baws\b': 'amazon web services',
            r'\bgcp\b': 'google cloud platform'
        }
        
        for pattern, replacement in tech_normalizations.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Clean up whitespace and special characters
        text = re.sub(r'[^\w\s\.\+\#\-/]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Keep technical terms and remove only basic stop words
        basic_stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                          'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being'}
        
        words = text.split()
        filtered_words = [word for word in words 
                         if word not in basic_stopwords or len(word) <= 3 or word.isdigit()]
        
        return " ".join(filtered_words)

    def extract_skills_with_context(self, text):
        """Extract skills with surrounding context for better matching"""
        skills_found = {}
        text_lower = text.lower()
        
        for skill_name, skill_data in self.skill_categories.items():
            terms = skill_data['terms']
            weight = skill_data['weight']
            category = skill_data['category']
            
            skill_matches = []
            for term in terms:
                # Look for the term with context
                pattern = rf'\b{re.escape(term)}\b'
                matches = list(re.finditer(pattern, text_lower))
                
                for match in matches:
                    start, end = match.span()
                    # Get context around the match
                    context_start = max(0, start - 50)
                    context_end = min(len(text_lower), end + 50)
                    context = text_lower[context_start:context_end]
                    
                    skill_matches.append({
                        'term': term,
                        'context': context,
                        'position': start
                    })
            
            if skill_matches:
                skills_found[skill_name] = {
                    'matches': skill_matches,
                    'count': len(skill_matches),
                    'weight': weight,
                    'category': category
                }
        
        return skills_found

    def calculate_semantic_similarity(self, cv_text, jd_text):
        """Advanced semantic similarity using multiple techniques"""
        try:
            # TF-IDF Cosine Similarity
            vectors = self.vectorizer.fit_transform([cv_text, jd_text])
            tfidf_similarity = cosine_similarity(vectors[0], vectors[1])[0][0]
            
            # Jaccard Similarity for exact matches
            cv_words = set(cv_text.split())
            jd_words = set(jd_text.split())
            jaccard_sim = len(cv_words.intersection(jd_words)) / len(cv_words.union(jd_words))
            
            # N-gram overlap
            def get_ngrams(text, n):
                words = text.split()
                return set([' '.join(words[i:i+n]) for i in range(len(words)-n+1)])
            
            cv_bigrams = get_ngrams(cv_text, 2)
            jd_bigrams = get_ngrams(jd_text, 2)
            cv_trigrams = get_ngrams(cv_text, 3)
            jd_trigrams = get_ngrams(jd_text, 3)
            
            bigram_sim = len(cv_bigrams.intersection(jd_bigrams)) / max(len(jd_bigrams), 1)
            trigram_sim = len(cv_trigrams.intersection(jd_trigrams)) / max(len(jd_trigrams), 1)
            
            # Weighted combination
            semantic_score = (
                tfidf_similarity * 0.4 +
                jaccard_sim * 0.3 +
                bigram_sim * 0.2 +
                trigram_sim * 0.1
            )
            
            return min(semantic_score, 1.0)
            
        except Exception as e:
            st.warning(f"Semantic similarity calculation error: {e}")
            return 0.0

    def calculate_skill_match_score(self, cv_skills, jd_skills):
        """Calculate weighted skill matching score"""
        total_weight = 0
        matched_weight = 0
        
        skill_breakdown = {
            'matched_skills': [],
            'missing_critical': [],
            'missing_nice_to_have': []
        }
        
        for skill_name, jd_skill_data in jd_skills.items():
            weight = jd_skill_data['weight']
            category = jd_skill_data['category']
            total_weight += weight
            
            if skill_name in cv_skills:
                # Calculate match strength based on frequency and context
                cv_count = cv_skills[skill_name]['count']
                jd_count = jd_skill_data['count']
                
                # Higher score for multiple mentions
                frequency_bonus = min(cv_count / max(jd_count, 1), 2.0)
                match_score = weight * frequency_bonus
                matched_weight += match_score
                
                skill_breakdown['matched_skills'].append({
                    'skill': skill_name,
                    'category': category,
                    'strength': frequency_bonus
                })
            else:
                if weight >= 2.0:  # Critical skills
                    skill_breakdown['missing_critical'].append({
                        'skill': skill_name,
                        'category': category,
                        'weight': weight
                    })
                else:  # Nice to have skills
                    skill_breakdown['missing_nice_to_have'].append({
                        'skill': skill_name,
                        'category': category,
                        'weight': weight
                    })
        
        skill_score = matched_weight / max(total_weight, 1)
        return skill_score, skill_breakdown

    def calculate_advanced_similarity(self, cv_text, job_description):
        """Enhanced similarity calculation with multiple factors"""
        if not cv_text or not job_description:
            return 0.0, {}
        
        try:
            # Preprocess texts
            cv_processed = self.advanced_text_preprocessing(cv_text)
            jd_processed = self.advanced_text_preprocessing(job_description)
            
            # Extract skills with context
            cv_skills = self.extract_skills_with_context(cv_text)
            jd_skills = self.extract_skills_with_context(job_description)
            
            # Calculate different similarity components
            semantic_score = self.calculate_semantic_similarity(cv_processed, jd_processed)
            skill_score, skill_breakdown = self.calculate_skill_match_score(cv_skills, jd_skills)
            
            # Experience and years calculation
            cv_years = self.extract_experience_years(cv_text)
            jd_years = self.extract_experience_years(job_description)
            experience_match = min(cv_years / max(jd_years, 1), 1.5) if jd_years > 0 else 1.0
            
            # Education and certification bonus
            education_bonus = self.calculate_education_bonus(cv_text, job_description)
            
            # Content quality factors
            cv_length = len(cv_text.split())
            quality_factor = min(cv_length / 200, 1.2)  # Bonus for comprehensive resumes
            
            # Weighted final score calculation
            base_score = (
                semantic_score * 0.25 +      # General semantic similarity
                skill_score * 0.45 +         # Technical skills (most important)
                experience_match * 0.15 +    # Experience level match
                education_bonus * 0.10 +     # Education/certification bonus
                quality_factor * 0.05        # Content quality
            )
            
            # Apply scaling for realistic scores
            if base_score >= 0.8:
                final_score = 85 + (base_score - 0.8) * 65  # 85-98%
            elif base_score >= 0.6:
                final_score = 70 + (base_score - 0.6) * 75  # 70-85%
            elif base_score >= 0.4:
                final_score = 50 + (base_score - 0.4) * 100 # 50-70%
            elif base_score >= 0.2:
                final_score = 30 + (base_score - 0.2) * 100 # 30-50%
            else:
                final_score = 15 + base_score * 75          # 15-30%
            
            # Additional bonuses
            if skill_score > 0.7:
                final_score += 5  # Strong technical skills bonus
            if len(skill_breakdown['missing_critical']) == 0:
                final_score += 8  # No missing critical skills bonus
            
            final_score = min(final_score, 98)  # Cap at 98%
            
            analysis_details = {
                'semantic_score': round(semantic_score * 100, 2),
                'skill_score': round(skill_score * 100, 2),
                'experience_match': round(experience_match * 100, 2),
                'education_bonus': round(education_bonus * 100, 2),
                'skill_breakdown': skill_breakdown,
                'cv_skills_count': len(cv_skills),
                'jd_skills_count': len(jd_skills),
                'cv_years': cv_years,
                'jd_years': jd_years
            }
            
            return round(final_score, 1), analysis_details
            
        except Exception as e:
            st.error(f"Error in similarity calculation: {e}")
            return 25.0, {}

    def extract_experience_years(self, text):
        """Extract years of experience from text"""
        patterns = [
            r'(\d+)[\+\-\s]*years?\s+(?:of\s+)?experience',
            r'(\d+)[\+\-\s]*yrs?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)[\+\-\s]*years?',
            r'(\d+)[\+\-\s]*years?\s+in',
            r'over\s+(\d+)\s+years?',
            r'more\s+than\s+(\d+)\s+years?'
        ]
        
        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            years.extend([int(match) for match in matches if match.isdigit()])
        
        return max(years) if years else 0

    def calculate_education_bonus(self, cv_text, jd_text):
        """Calculate bonus for education and certifications"""
        cv_lower = cv_text.lower()
        jd_lower = jd_text.lower()
        
        education_terms = ['bachelor', 'master', 'phd', 'degree', 'computer science', 
                          'software engineering', 'information technology']
        cert_terms = ['certified', 'certification', 'aws certified', 'google certified',
                     'microsoft certified', 'mongodb certified']
        
        bonus = 0.0
        
        # Education matching
        for term in education_terms:
            if term in jd_lower and term in cv_lower:
                bonus += 0.1
        
        # Certification matching
        for term in cert_terms:
            if term in jd_lower and term in cv_lower:
                bonus += 0.15
        
        return min(bonus, 0.5)  # Cap at 50% bonus

    def generate_comprehensive_feedback(self, cv_text, job_description, score, analysis_details):
        """Generate detailed, actionable feedback"""
        try:
            feedback_sections = []
            
            # Overall assessment
            if score >= 85:
                feedback_sections.append("🎉 **Excellent Match!** Your resume strongly aligns with this MERN stack position.")
            elif score >= 70:
                feedback_sections.append("🚀 **Strong Candidate!** Your profile shows great potential for this role.")
            elif score >= 55:
                feedback_sections.append("👍 **Good Foundation** with clear opportunities for improvement.")
            elif score >= 40:
                feedback_sections.append("📈 **Decent Match** but requires targeted enhancements.")
            else:
                feedback_sections.append("⚠️ **Significant Gap** - consider major resume optimization.")
            
            # Skill analysis
            skill_breakdown = analysis_details.get('skill_breakdown', {})
            
            if skill_breakdown.get('matched_skills'):
                matched = skill_breakdown['matched_skills'][:6]  # Top 6
                matched_list = [f"**{s['skill'].title()}** ({s['category']})" for s in matched]
                feedback_sections.append(f"✅ **Strong Technical Skills:** {', '.join(matched_list)}")
            
            # Critical missing skills
            if skill_breakdown.get('missing_critical'):
                missing_critical = skill_breakdown['missing_critical'][:4]
                critical_list = [f"**{s['skill'].title()}**" for s in missing_critical]
                feedback_sections.append(f"🚨 **Critical Missing Skills:** {', '.join(critical_list)}")
                feedback_sections.append("💡 *Focus on these first - they're essential for MERN stack roles.*")
            
            # Nice to have skills
            if skill_breakdown.get('missing_nice_to_have'):
                nice_to_have = skill_breakdown['missing_nice_to_have'][:3]
                nice_list = [f"**{s['skill'].title()}**" for s in nice_to_have]
                feedback_sections.append(f"📋 **Additional Skills to Consider:** {', '.join(nice_list)}")
            
            # Experience feedback
            cv_years = analysis_details.get('cv_years', 0)
            jd_years = analysis_details.get('jd_years', 0)
            if jd_years > 0:
                if cv_years >= jd_years:
                    feedback_sections.append(f"✅ **Experience Match:** {cv_years} years (meets {jd_years} year requirement)")
                else:
                    feedback_sections.append(f"⏰ **Experience Gap:** {cv_years} years (needs {jd_years} years)")
            
            # Detailed scores
            feedback_sections.append("## 📊 **Detailed Analysis**")
            feedback_sections.append(f"- **Technical Skills Match:** {analysis_details.get('skill_score', 0)}%")
            feedback_sections.append(f"- **Content Relevance:** {analysis_details.get('semantic_score', 0)}%") 
            feedback_sections.append(f"- **Experience Alignment:** {analysis_details.get('experience_match', 0)}%")
            
            # Actionable recommendations
            feedback_sections.append("## 🎯 **Recommendations**")
            
            if score < 70:
                recommendations = [
                    "🔧 Add specific MERN stack project examples",
                    "📝 Include more technical keywords from the job description",
                    "💼 Highlight relevant work experience with quantified achievements",
                    "🏆 Add any relevant certifications or courses"
                ]
                feedback_sections.append("\n".join(recommendations))
            else:
                recommendations = [
                    "✨ Fine-tune keyword optimization",
                    "📈 Add metrics and quantified achievements",
                    "🔍 Ensure all relevant projects are highlighted"
                ]
                feedback_sections.append("\n".join(recommendations))
            
            return "\n\n".join(feedback_sections)
            
        except Exception as e:
            return self.generate_fallback_feedback(score)

    def generate_fallback_feedback(self, score):
        """Fallback feedback when detailed analysis fails"""
        if score >= 75:
            return "🎉 **Excellent Match!** Your resume shows strong alignment with MERN stack requirements."
        elif score >= 60:
            return "👍 **Good Match** with some optimization opportunities for better alignment."
        elif score >= 45:
            return "📊 **Fair Match** - focus on highlighting MERN stack experience and skills."
        else:
            return "📈 **Needs Improvement** - consider adding more relevant MERN stack experience and keywords."

    def process_resume(self, cv_file, job_description):
        """Main processing function with enhanced error handling"""
        try:
            # Extract and validate text
            raw_cv_text = self.extract_text(cv_file)
            
            if not raw_cv_text.strip():
                raise ValueError("Could not extract meaningful text from the resume file")
            
            if len(raw_cv_text.split()) < 50:
                st.warning("⚠️ Resume seems quite short. Consider adding more details for better analysis.")
            
            # Perform analysis
            score, analysis_details = self.calculate_advanced_similarity(raw_cv_text, job_description)
            feedback = self.generate_comprehensive_feedback(raw_cv_text, job_description, score, analysis_details)
            
            # Determine match level
            if score >= 80:
                match_level = "Excellent Match"
                match_color = "🟢"
            elif score >= 65:
                match_level = "Very Good Match"
                match_color = "🟡"
            elif score >= 50:
                match_level = "Good Match"
                match_color = "🟡"
            elif score >= 35:
                match_level = "Fair Match"
                match_color = "🟠"
            else:
                match_level = "Needs Improvement"
                match_color = "🔴"
            
            return {
                "score": score,
                "feedback": feedback,
                "match_level": match_level,
                "match_color": match_color,
                "cv_length": len(raw_cv_text.split()),
                "jd_length": len(job_description.split()),
                "analysis_details": analysis_details
            }
            
        except Exception as e:
            raise ValueError(f"Error processing resume: {str(e)}")

# --- Enhanced Streamlit Interface ---

st.set_page_config(
    page_title="MERN Stack Resume Matcher", 
    page_icon="⚛️", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .stProgress .st-bo {
        background-color: #e6f3ff;
    }
</style>
""", unsafe_allow_html=True)

st.title("⚛️ MERN Stack Resume Matcher")
st.markdown("**Advanced AI-powered resume analysis specifically optimized for MERN stack positions**")
st.markdown("---")

# Initialize the enhanced matcher
if 'enhanced_matcher' not in st.session_state:
    st.session_state.enhanced_matcher = EnhancedMERNResumeMatcher()

matcher = st.session_state.enhanced_matcher

# Main layout
col1, col2 = st.columns([1, 1])

with col1:
    st.header("📄 Upload Your Resume")
    cv_file = st.file_uploader(
        "Choose your resume file", 
        type=["pdf", "docx", "txt"],
        help="Supported formats: PDF, DOCX, TXT (Max size: 10MB)"
    )
    
    if cv_file:
        st.success(f"✅ Resume uploaded: {cv_file.name}")
        file_size = len(cv_file.getvalue()) / 1024
        st.info(f"📊 File size: {file_size:.1f} KB")

with col2:
    st.header("💼 Job Description")
    input_method = st.radio(
        "Input method:",
        ("📝 Paste Description", "📁 Upload File"),
        horizontal=True
    )

job_description = ""

if input_method == "📝 Paste Description":
    job_description = st.text_area(
        "Paste the MERN stack job description", 
        height=200,
        placeholder="Paste the complete job description including requirements, responsibilities, and preferred qualifications..."
    )
    if job_description:
        word_count = len(job_description.split())
        st.info(f"📊 Word count: {word_count}")
        
        if word_count < 100:
            st.warning("⚠️ Job description seems short. More details will improve analysis accuracy.")

else:
    job_file = st.file_uploader(
        "Upload job description file", 
        type=["pdf", "docx", "txt"],
        help="Upload a file containing the complete job description"
    )
    if job_file:
        try:
            job_description = matcher.extract_text(job_file)
            st.success("✅ Job description loaded successfully!")
            word_count = len(job_description.split())
            st.info(f"📊 Word count: {word_count}")
        except Exception as e:
            st.error(f"❌ Error loading job description: {e}")

# Analysis section
st.markdown("---")
st.header("🔍 Advanced Analysis")

if st.button("🚀 Analyze MERN Stack Match", type="primary", use_container_width=True):
    if not cv_file:
        st.error("❌ Please upload your resume file.")
    elif not job_description:
        st.error("❌ Please provide the job description.")
    elif len(job_description.split()) < 20:
        st.error("❌ Job description is too short for meaningful analysis.")
    else:
        try:
            with st.spinner("🔄 Analyzing your MERN stack resume... This may take a moment."):
                result = matcher.process_resume(cv_file, job_description)
            
            # Display results with enhanced formatting
            score = result['score']
            match_level = result['match_level']
            match_color = result['match_color']
            analysis_details = result.get('analysis_details', {})
            
            st.markdown("## 📊 Analysis Results")
            
            # Enhanced metrics display
            met_col1, met_col2, met_col3, met_col4 = st.columns(4)
            
            with met_col1:
                st.metric(
                    label="Overall Match Score", 
                    value=f"{score}%",
                    help="Comprehensive compatibility score for MERN stack position"
                )
            
            with met_col2:
                st.metric(
                    label="Technical Skills", 
                    value=f"{analysis_details.get('skill_score', 0):.1f}%",
                    help="Technical skills alignment with job requirements"
                )
            
            with met_col3:
                st.metric(
                    label="Content Relevance", 
                    value=f"{analysis_details.get('semantic_score', 0):.1f}%",
                    help="Semantic similarity between resume and job description"
                )
            
            with met_col4:
                st.metric(
                    label="Experience Match", 
                    value=f"{analysis_details.get('experience_match', 0):.1f}%",
                    help="Experience level alignment with requirements"
                )
            
            # Enhanced progress visualization
            st.markdown("### 📈 Score Breakdown")
            
            # Main score progress bar
            progress_col1, progress_col2 = st.columns([3, 1])
            with progress_col1:
                st.progress(score / 100)
            with progress_col2:
                st.markdown(f"**{match_color} {score}%**")
            
            st.markdown(f"**Match Level:** {match_level}")
            
            # Detailed component scores
            if analysis_details:
                st.markdown("#### Component Scores")
                comp_col1, comp_col2 = st.columns(2)
                
                with comp_col1:
                    st.write("**Technical Skills**")
                    st.progress(analysis_details.get('skill_score', 0) / 100)
                    st.write("**Content Relevance**")
                    st.progress(analysis_details.get('semantic_score', 0) / 100)
                
                with comp_col2:
                    st.write("**Experience Alignment**")
                    st.progress(analysis_details.get('experience_match', 0) / 100)
                    st.write("**Education/Certifications**")
                    st.progress(analysis_details.get('education_bonus', 0) / 100)
            
            # Detailed feedback
            st.markdown("---")
            st.markdown("### 💡 Comprehensive Feedback")
            st.markdown(result['feedback'])
            
            # Skills breakdown visualization
            if analysis_details.get('skill_breakdown'):
                st.markdown("---")
                st.markdown("### 🛠️ Skills Analysis")
                
                skill_breakdown = analysis_details['skill_breakdown']
                
                # Create tabs for different skill categories
                skill_tab1, skill_tab2, skill_tab3 = st.tabs(["✅ Matched Skills", "🚨 Missing Critical", "📋 Nice to Have"])
                
                with skill_tab1:
                    if skill_breakdown.get('matched_skills'):
                        st.success(f"Found {len(skill_breakdown['matched_skills'])} matching technical skills!")
                        
                        # Group by category
                        categories = {}
                        for skill in skill_breakdown['matched_skills']:
                            cat = skill['category']
                            if cat not in categories:
                                categories[cat] = []
                            categories[cat].append(skill)
                        
                        for category, skills in categories.items():
                            with st.expander(f"📂 {category} ({len(skills)} skills)"):
                                for skill in skills:
                                    strength_bar = "🟢" * min(int(skill['strength']), 5)
                                    st.write(f"• **{skill['skill'].replace('_', ' ').title()}** {strength_bar}")
                    else:
                        st.warning("No matching technical skills found. Consider adding MERN stack technologies to your resume.")
                
                with skill_tab2:
                    if skill_breakdown.get('missing_critical'):
                        st.error(f"⚠️ {len(skill_breakdown['missing_critical'])} critical skills missing!")
                        st.markdown("**These skills are essential for MERN stack roles:**")
                        
                        for skill in skill_breakdown['missing_critical']:
                            st.write(f"🔴 **{skill['skill'].replace('_', ' ').title()}** ({skill['category']})")
                        
                        st.info("💡 **Tip:** Add projects or experience with these technologies to significantly improve your match score.")
                    else:
                        st.success("🎉 Great! No critical skills are missing from your resume.")
                
                with skill_tab3:
                    if skill_breakdown.get('missing_nice_to_have'):
                        st.info(f"📋 {len(skill_breakdown['missing_nice_to_have'])} additional skills could strengthen your profile:")
                        
                        for skill in skill_breakdown['missing_nice_to_have']:
                            st.write(f"🟡 **{skill['skill'].replace('_', ' ').title()}** ({skill['category']})")
                        
                        st.info("💡 **Tip:** These skills aren't critical but could give you an edge over other candidates.")
                    else:
                        st.success("✨ You have excellent coverage of desired skills!")
            
            # Experience analysis
            if analysis_details.get('cv_years') is not None and analysis_details.get('jd_years') is not None:
                st.markdown("---")
                st.markdown("### 📅 Experience Analysis")
                
                cv_years = analysis_details['cv_years']
                jd_years = analysis_details['jd_years']
                
                exp_col1, exp_col2 = st.columns(2)
                
                with exp_col1:
                    st.metric("Your Experience", f"{cv_years} years")
                
                with exp_col2:
                    st.metric("Required Experience", f"{jd_years} years")
                
                if cv_years >= jd_years:
                    st.success(f"✅ You meet the experience requirement! ({cv_years} >= {jd_years} years)")
                elif cv_years > 0:
                    gap = jd_years - cv_years
                    st.warning(f"⏰ Experience gap: {gap} year(s). Consider highlighting relevant projects or internships.")
                else:
                    st.info("💡 No specific experience years found in your resume. Consider adding your years of experience.")
            
            # Action items and next steps
            st.markdown("---")
            st.markdown("### 🎯 Action Items")
            
            if score >= 85:
                st.balloons()
                st.success("🏆 **Outstanding!** Your resume is excellently positioned for this MERN stack role.")
                st.markdown("""
                **Next Steps:**
                - Review the job description for any specific project requirements
                - Prepare examples of your MERN stack projects for interviews
                - Consider reaching out to the hiring team
                """)
            elif score >= 70:
                st.success("🚀 **Strong Candidate!** You're well-positioned for this role.")
                st.markdown("""
                **Optimization Tips:**
                - Add 1-2 more relevant projects to your portfolio
                - Include specific metrics and achievements
                - Ensure all required technologies are mentioned
                """)
            elif score >= 55:
                st.info("📈 **Good Potential** with room for targeted improvements.")
                st.markdown("""
                **Priority Actions:**
                - Address any missing critical skills through projects or courses
                - Expand your technical skills section
                - Add more detailed project descriptions
                """)
            else:
                st.warning("⚠️ **Significant Enhancement Needed**")
                st.markdown("""
                **Focus Areas:**
                - Build projects using the MERN stack
                - Complete relevant online courses or certifications
                - Restructure resume to highlight technical skills prominently
                - Consider gaining more hands-on experience with missing technologies
                """)
            
            # Export results option
            st.markdown("---")
            if st.button("📥 Export Analysis Report"):
                report_data = {
                    "overall_score": score,
                    "match_level": match_level,
                    "analysis_date": st.session_state.get('analysis_date', 'Unknown'),
                    "detailed_scores": analysis_details,
                    "recommendations": result['feedback']
                }
                
                st.download_button(
                    label="Download JSON Report",
                    data=json.dumps(report_data, indent=2),
                    file_name=f"resume_analysis_{score:.0f}pct.json",
                    mime="application/json"
                )
                
        except Exception as e:
            st.error(f"❌ Error during analysis: {e}")
            st.info("💡 **Troubleshooting:**")
            st.info("• Ensure your resume file is not corrupted")
            st.info("• Try a different file format (PDF recommended)")
            st.info("• Check that the job description contains sufficient detail")

# Enhanced sidebar with MERN-specific tips
with st.sidebar:
    st.markdown("## ⚛️ MERN Stack Tips")
    st.markdown("""
    **🔥 Critical Technologies:**
    - **MongoDB** - Document database
    - **Express.js** - Web framework  
    - **React** - Frontend library
    - **Node.js** - Runtime environment
    
    **💡 Score Boosters:**
    - Include specific project examples
    - Mention relevant npm packages
    - Show full-stack project experience
    - Add deployment experience (Heroku, Netlify, AWS)
    
    **📈 Optimization Guide:**
    - **85-98%:** Excellent - Ready to apply! ✅
    - **70-84%:** Very Good - Minor tweaks needed 👍
    - **55-69%:** Good - Address missing skills 📊
    - **40-54%:** Fair - Need significant improvements ⚠️
    - **Below 40%:** Needs major rework 🔄
    
    **🚀 Pro Tips:**
    - Quantify your achievements with metrics
    - Include GitHub portfolio links
    - Mention Agile/Scrum experience
    - Show responsive design skills
    - Add testing experience (Jest, Cypress)
    """)
    
    st.markdown("---")
    
    st.markdown("## 📚 Learning Resources")
    st.markdown("""
    **Free MERN Learning:**
    - [React Documentation](https://reactjs.org)
    - [Node.js Guides](https://nodejs.org)
    - [MongoDB Tutorial](https://docs.mongodb.com)
    - [Express.js Guide](https://expressjs.com)
    
    **Project Ideas:**
    - E-commerce application
    - Social media dashboard
    - Task management app
    - Real-time chat application
    - Blog platform with CMS
    """)
    
    st.markdown("---")
    st.markdown("**⚛️ Built for MERN Stack Success**")
    st.markdown("*Enhanced with advanced NLP and skill matching*")